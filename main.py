import openpyxl
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from tkinter import Tk, filedialog, Label, Button, Entry, StringVar, OptionMenu, Frame, Spinbox


def format_headers(headers: list[str]) -> list[str]:
    """
    Formats the headers to have capitalized first letters on each word.
    :param headers: List of the respective headers
    """
    headers = [header.title() for header in headers if isinstance(header, str)]
    return headers


def format_values(values: list[float]) -> list[float]:
    """
    Formats the values to be rounded based on the common rounding rule and appearing as whole percentages.
    :param values: List of the respective values
    """
    values = [round(value * 100) if isinstance(value, (int, float)) else value for value in values]
    return values


def read_excel(file_path: str) -> dict[str, dict[str, list[str]]]:
    """
    Reads an Excel file and extracts headers from the first column and values from the second column
    for all sheets. Returns a dictionary with sheet names as keys and extracted data as values.
    :param file_path: Absolute path to file
    """
    workbook = openpyxl.load_workbook(file_path)
    data = {}

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        headers = [row[0] for row in sheet.iter_rows(
            min_row=1, max_row=sheet.max_row, min_col=1, max_col=1, values_only=True
        )]
        values = [row[0] for row in sheet.iter_rows(
            min_row=1, max_row=sheet.max_row, min_col=2, max_col=2, values_only=True
        )]

        try:
            start_index = values.index(1) + 1
            filtered_headers = headers[start_index:]
            filtered_headers = format_headers(filtered_headers)
            filtered_values = values[start_index:]
            filtered_values = format_values(filtered_values)
        except ValueError:
            filtered_headers = []
            filtered_values = []

        data[sheet_name] = {"headers": filtered_headers, "values": filtered_values}

    return data


def get_questions(file_path: str) -> list[str]:
    """
    Reads an Excel file and extracts the questions on each sheet
    :param file_path: Absolute path to file
    """
    workbook = openpyxl.load_workbook(file_path)
    questions = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        headers = [row[0] for row in sheet.iter_rows(
            min_row=1, max_row=sheet.max_row, min_col=1, max_col=1, values_only=True
        )]
        questions.append(headers[2])
    return questions


def style_table(table, args: dict[str, str]) -> None:
    """
    Styles the table based on the provided arguments.
    :param table: The table to be styled.
    :param args: Dictionary containing report options such as total position, font type, font size, ordering, etc.
    """

    # for row in table.rows:
    #     for i, cell in enumerate(row.cells):
    #         if i == 0:
    #             cell.width = Inches(1)
    #         else:
    #             cell.width = Inches(1.5)

    # Remove gridlines
    # tbl = table._tbl
    # for tblBorders in tbl.xpath(".//w:tblBorders"):
    #     tblBorders.getparent().remove(tblBorders)

    if args.get("font_type"):
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.name = args.get("font_type")
                cell.paragraphs[0].runs[0].font.size = Pt(int(args.get("font_size")))
    if args.get("text_type") == "All Caps":
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = paragraph.text.upper()
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if "total" in paragraph.text.lower():
                    for run in paragraph.runs:
                        if args.get("total_position") == "inline":
                            run.bold = True
                            run.underline = WD_UNDERLINE.SINGLE
                        else:
                            run.bold = True
                            run.italic = True
                            run.text = run.text.upper()

                    connected_cell = row.cells[i - 1] if args.get("header_side") == "Right" else row.cells[i + 1]
                    for connected_paragraph in connected_cell.paragraphs:
                        for connected_run in connected_paragraph.runs:
                            if args.get("total_position") == "inline":
                                connected_run.bold = True
                                connected_run.underline = WD_UNDERLINE.SINGLE
                            else:
                                connected_run.bold = True
                                connected_run.italic = True
                                connected_run.text = connected_run.text.upper()

    # TODO: Formatting totals on top, bottom or inline (selector)
    return table


def add_percentages_to_values(value: str) -> str:
    percentage_dict = ["--", "*"]
    if value in percentage_dict:
        return str(value)
    elif value is not None:
        return str(value) + "%"
    else:
        return ""

def write_doc(data: dict[str, dict[str, list[str]]], questions: list[str], output_path: str,
              args: dict[str, str]) -> None:
    """
    Writes dictionaries of headers and values into tables in a .docx file.
    :param data: Dictionary with sheet names as keys and dictionaries of headers and values as values.
    :param questions: List of question associated with each sheet.
    :param output_path: Path to save the generated Word document.
    :param args: Dictionary containing report options such as total position, font type, font size, ordering, etc.
    """

    document = Document()
    i = 0
    other_dict = ["other", "unsure", "refused", "no opinion"]
    for sheet_name, content in data.items():
        document.add_heading(sheet_name, level=1)
        document.add_paragraph(questions[i])
        if args.get("ordering") == "Vertical":
            table = document.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for header, value in zip(content["headers"], content["values"]):
                if header and (header.lower() in other_dict or "total" in header.lower()):
                    blank_row = table.add_row()
                    blank_row.cells[0].text = ""
                    blank_row.cells[1].text = ""
                row = table.add_row()
                if args.get("header_side") == "Right":
                    value_cell = row.cells[0]
                    value_cell.text = add_percentages_to_values(value)
                    value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    row.cells[1].text = str(header) if header is not None else ""
                else:
                    value_cell = row.cells[1]
                    value_cell.text = add_percentages_to_values(value)
                    value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    row.cells[0].text = str(header) if header is not None else ""

        elif args.get("ordering") == "Horizontal":
            num_cols = len(content["headers"])
            table = document.add_table(rows=2, cols=num_cols)
            table.style = 'Table Grid'
            for col, header in enumerate(content["headers"]):
                table.cell(0, col).text = str(header) if header is not None else ""
            for col, value in enumerate(content["values"]):
                value_cell = table.cell(1, col)
                value_cell.text = str(value) + "%" if value is not None else ""
                value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        style_table(table, args)
        i += 1

    document.save(output_path)


def run_report(file_path: str, args: dict[str, str]) -> None:
    """
    Runs the report generation process by reading the Excel file, extracting data,
    and writing it to a Word document.
    :param file_path: Absolute path to the Excel file.
    :param args: Dictionary containing report options such as total position, font type, font size, and custom title.
    """
    suffix = "_v" if args.get("ordering") == "Vertical" else "_h"
    excel_data = read_excel(file_path)
    if file_path.lower().endswith('.xlsx'):
        output_file_path = file_path[:-5] + suffix + ".docx"
    elif file_path.lower().endswith('.xls'):
        output_file_path = file_path[:-4] + suffix + ".docx"
    else:
        exit("Invalid input file")
    questions = get_questions(file_path)
    write_doc(excel_data, questions, output_file_path, args)
    print(f"Report written to {output_file_path}")


def open_gui() -> None:
    """
    Opens a GUI for selecting an Excel file and configuring report options.
    The GUI allows the user to select a file, choose the position of totals, font type, font size,
    text type, and run the report generation process.
    The GUI uses Tkinter for the interface.
    """
    root = Tk()
    root.title("Report Writer")
    root.geometry("350x550")

    file_frame = Frame(root)
    file_frame.pack(pady=10)
    file_label = Label(file_frame, text="No file selected")
    file_label.pack(side="left")

    def select_file():
        path = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        file_var.set(path)
        file_label.config(text=os.path.basename(path) if path else "No file selected")

    file_var = StringVar()
    Button(
        file_frame,
        text="Browse",
        command=select_file,
        activebackground="#cccccc",
        activeforeground="#000000"
    ).pack(side="left", padx=5)

    common_fonts = ["Arial", "Calibri", "Times New Roman", "Verdana", "Courier New", "Georgia", "Tahoma", "Helvetica"]
    Label(root, text="Font Type:").pack()
    font_type_var = StringVar(value="Tahoma")
    font_dropdown = OptionMenu(root, font_type_var, *common_fonts)
    font_dropdown.pack()

    Label(root, text="\nOr enter a custom font:").pack()
    custom_font_var = StringVar()
    Entry(root, textvariable=custom_font_var).pack()

    font_warning = Label(root, text="", fg="red")
    font_warning.pack()

    def on_font_entry(*args):
        if custom_font_var.get():
            font_warning.config(text="Warning: Custom font may not be available in Word.")
            font_type_var.set(custom_font_var.get())
        else:
            font_warning.config(text="")

    custom_font_var.trace_add("write", on_font_entry)

    Label(root, text="Font Size:").pack()
    font_size_var = StringVar(value="9")
    Spinbox(
        root,
        values=tuple(str(i) for i in range(6, 25)),
        textvariable=font_size_var,
        width=5,
        state="readonly"
    ).pack()
    font_size_var.set("9")

    total_position_var = StringVar(value="Top")
    Label(root, text="\nTotals Position: (Not implemented)").pack()
    OptionMenu(root, total_position_var, "Top", "Bottom", "Inline").pack()

    Label(root, text="\nText Type:").pack()
    text_type_var = StringVar(value="Title")
    OptionMenu(root, text_type_var, "Title", "All Caps").pack()

    Label(root, text="\nHeader Side:").pack()
    header_side_var = StringVar(value="Right")
    OptionMenu(root, header_side_var, "Right", "Left").pack()

    Label(root, text="\nOrdering:").pack()
    ordering_var = StringVar(value="Vertical")
    OptionMenu(root, ordering_var, "Vertical", "Horizontal").pack()

    def on_run():
        if not file_var.get():
            file_label.config(text="Please select a file!")
            return
        args = {
            "total_position": total_position_var.get(),
            "font_type": font_type_var.get(),
            "font_size": font_size_var.get(),
            "text_type": text_type_var.get(),
            "header_side": header_side_var.get(),
            "ordering": ordering_var.get()
        }
        run_report(file_var.get(), args)
        root.quit()

    Button(
        root,
        text="Generate Report",
        command=on_run,
        activebackground="#cccccc",
        activeforeground="#000000"
    ).pack(pady=10)
    root.mainloop()


if __name__ == "__main__":
    open_gui()
