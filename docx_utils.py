from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from helpers import add_percentages_to_values, move_totals

ARGS = {}


def initialize_args(args: dict[str, str]) -> None:
    """
    Initializes the global ARGS dictionary with values from the provided args.
    :param args: Dictionary containing report options.
    """
    global ARGS
    ARGS = {
        "total_position": args.get("total_position", ""),
        "ordering": args.get("ordering", ""),
        "header_side": args.get("header_side", ""),
        "gridlines": args.get("gridlines", False),
        "font_size": int(args.get("font_size", 12)),
        "font_type": args.get("font_type", ""),
        "text_type": args.get("text_type", ""),
        "margin": float(args.get("margin", "")),
        "extra_cols": args.get("extra_columns_flag", False)
    }


# Thanks Copilot <3
def remove_table_borders(table):
    """
    Removes gridlines (borders) from a table in a Word document.
    :param table: The table object from python-docx.
    """
    tbl = table._tbl  # Access the table's XML
    tbl_pr = tbl.tblPr  # Access the table properties

    # Check if <w:tblBorders> exists
    tbl_borders = tbl_pr.xpath(".//w:tblBorders")
    if tbl_borders:
        # Remove existing <w:tblBorders>
        for border in tbl_borders:
            border.getparent().remove(border)
    else:
        # Add <w:tblBorders> with all borders set to "none"
        borders_xml = (
            '<w:tblBorders {}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        ).format(nsdecls('w'))
        tbl_pr.append(parse_xml(borders_xml))


def set_table_margin(table, margin: Inches) -> None:
    """
    Explicitly sets the margin when the headers are on the left of a vertical table
    :param table: The table object from python-docx.
    :param margin: The margin value in inches.
    """
    tbl = table._tbl  # Access the table's XML
    margin_twips = int(margin * 1440)  # Convert inches to twips (1 inch = 1440 twips)
    tbl_pr = tbl.tblPr  # Access the table properties

    # Add or modify the <w:tblInd> property
    tbl_ind_xml = f'<w:tblInd {nsdecls("w")} w:w="{margin_twips}" w:type="dxa"/>'
    tbl_pr.append(parse_xml(tbl_ind_xml))


def style_table(table) -> None:
    """
    Styles the table based on the provided arguments.
    :param table: The table to be styled.
    """

    if ARGS["ordering"] == "Vertical":
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if ARGS["header_side"] == "Right" and i == 0:
                    cell.width = Inches(ARGS["margin"])
                elif ARGS["header_side"] == "Left" and i != 0:
                    cell.width = Inches(ARGS["margin"])

        if ARGS["header_side"] == "Left":
            set_table_margin(table, ARGS["margin"])

    if not ARGS["gridlines"]:
        remove_table_borders(table)

    if ARGS["text_type"] == "All Caps":
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = paragraph.text.upper()
                    
    if ARGS["font_type"]:
        for row in table.rows:
            for cell in row.cells:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.name = ARGS["font_type"]
                    cell.paragraphs[0].runs[0].font.size = Pt(ARGS["font_size"])

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if "total" in paragraph.text.lower():
                    for run in paragraph.runs:
                        if ARGS["total_position"] == "Inline":
                            run.bold = True
                            run.underline = WD_UNDERLINE.SINGLE
                        else:
                            run.bold = True
                            run.italic = True
                            run.text = run.text.upper()

                    connected_cell = row.cells[i - 1] \
                        if (ARGS["header_side"] == "Right" and ARGS["ordering"] == "Vertical") else row.cells[i + 1]
                    for connected_paragraph in connected_cell.paragraphs:
                        for connected_run in connected_paragraph.runs:
                            if ARGS["total_position"] == "Inline":
                                connected_run.bold = True
                                connected_run.underline = WD_UNDERLINE.SINGLE
                            else:
                                connected_run.bold = True
                                connected_run.italic = True
                                connected_run.text = connected_run.text.upper()

    return table


def gen_vert_table(document: Document, content: dict[str, list[str]]):
    """
    Generates a vertical table in a Word document based on the provided content and styles it according to ARGS.

    :param document: The Word document object where the table will be added.
    :param content: A dictionary containing "headers" and "values" as keys with their respective lists.
    :return: The generated and styled vertical table object.
    """

    num_cols = 2 + len(content.get("values")) - 1
    table = document.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'
    for header, values in zip(content["headers"], zip(*content["values"])):
        row = table.add_row()
        if ARGS["header_side"] == "Right":
            for col, value in enumerate(values):
                value_cell = row.cells[col]
                value_cell.text = add_percentages_to_values(value)
                value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row.cells[num_cols - 1].text = str(header) if header is not None else ""
        else:
            for col, value in enumerate(values):
                value_cell = row.cells[col + 1]
                value_cell.text = add_percentages_to_values(value)
                value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            row.cells[0].text = str(header) if header is not None else ""

    return table


def gen_horiz_table(document: Document, content: dict[str, list[str]]):
    """
    Generates a horizontal table in a Word document based on the provided content.

    :param document: The Word document object where the table will be added.
    :param content: A dictionary containing "headers" and "values" as keys with their respective lists.
    :return: The generated horizontal table object.
    """
    num_cols = len(content["headers"])
    num_rows = 2 + len(content.get("values")) - 1
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    for col, header in enumerate(content["headers"]):
        table.cell(0, col).text = str(header) if header is not None else "  "
    for row, values in enumerate(content["values"]):
        for col, value in enumerate(values):
            value_cell = table.cell(row+1, col)
            value_cell.text = add_percentages_to_values(value)
            value_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return table


def write_doc(data: dict[str, dict[str, list[str]]], pre_data: list[list[str]], output_path: str,
              args: dict[str, str]) -> None:
    """
    Writes dictionaries of headers and values into tables in a .docx file.
    :param data: Dictionary with sheet names as keys and dictionaries of headers and values as values.
    :param pre_data: List of pre-data content associated with each sheet.
    :param output_path: Path to save the generated Word document.
    :param args: Dictionary containing report options such as total position, font type, font size, ordering, etc.
    """
    initialize_args(args)

    document = Document()
    i = 0

    for sheet_name, content in data.items():
        document.add_heading(sheet_name, level=1)
        for question_data in pre_data[i]:
            document.add_paragraph(question_data)
        for i, values in enumerate(content["values"]):
            if ARGS["total_position"] == "Bottom":
                content["headers"], content["values"][i] = move_totals(content["headers"], values, "Bottom")
            elif ARGS["total_position"] == "Top":
                content["headers"], content["values"][i] = move_totals(content["headers"], values, "Top")

        if ARGS["ordering"] == "Vertical":
            table = gen_vert_table(document, content)
        elif ARGS["ordering"] == "Horizontal":
            table = gen_horiz_table(document, content)
        else:
            table_v = gen_vert_table(document, content)
            ARGS["ordering"] = "Vertical"
            style_table(table_v)
            document.add_paragraph("\n")
            for i, values in enumerate(content["values"]):
                if ARGS["total_position"] != "Inline":
                    content["headers"], content["values"][i] = move_totals(content["headers"], values, "Top")
            table_h = gen_horiz_table(document, content)
            ARGS["ordering"] = "Horizontal"
            style_table(table_h)
            document.add_page_break()
            ARGS["ordering"] = "Both"

        i += 1
        if ARGS["ordering"] == "Vertical" or ARGS["ordering"] == "Horizontal":
            style_table(table)
            if i % 2 == 0:
                document.add_page_break()

    document.save(output_path)
